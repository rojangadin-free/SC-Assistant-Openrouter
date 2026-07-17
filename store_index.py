import os
import re
import time
import io
import json
import logging
import tempfile
from typing import List, Optional, Iterable, Tuple

import fitz  # PyMuPDF
import pdfplumber
import docx
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from PIL import Image, ImageOps, ImageFilter, ImageEnhance
from dotenv import load_dotenv
from pinecone import Pinecone, ServerlessSpec
from langchain_community.retrievers import PineconeHybridSearchRetriever
from langchain_core.documents import Document
from langchain_community.document_loaders import (
    TextLoader,
    UnstructuredMarkdownLoader,
    CSVLoader,
)
from pinecone_text.sparse import BM25Encoder

# --- NEW IMPORTS FOR PYDANTIC EXTRACTION ---
from pydantic import BaseModel, Field
from langchain_core.messages import HumanMessage

from src.helper import (
    smart_chunking,
    get_local_embeddings,
    generate_image_caption,
    clean_text,
    generate_chunk_id,
    is_low_value_chunk,
    get_vision_client,     # Added for structured output
    encode_image           # Added for structured output
)

from aws.s3 import upload_file_to_s3

logger = logging.getLogger(__name__)

load_dotenv()
PINECONE_API_KEY = os.environ.get("PINECONE_API_KEY")
if not PINECONE_API_KEY:
    raise RuntimeError("PINECONE_API_KEY not set in environment.")

BM25_VALUES_FILE = "bm25_values.json"


# ============================================================
# PYDANTIC STRUCTURED OUTPUT MODELS (For Tables)
# ============================================================
class TableRow(BaseModel):
    column_1: str = Field(description="Data in the first column, or empty string")
    column_2: str = Field(description="Data in the second column, or empty string")
    column_3: str = Field(description="Data in the third column, or empty string")
    column_4: str = Field(description="Data in the fourth column, or empty string")
    column_5: str = Field(description="Data in the fifth column, or empty string")
    column_6: str = Field(description="Data in the sixth column, or empty string")
    column_7: str = Field(description="Data in the seventh column, or empty string")

class ExtractedTable(BaseModel):
    rows: List[TableRow] = Field(description="All rows extracted from the table in exact order")


# ============================================================
# METADATA SANITIZER
# ============================================================
def clean_metadata(doc: Document) -> Document:
    cleaned = {}
    for k, v in (doc.metadata or {}).items():
        if k in {"image_base64", "base64_image", "orig_elements", "coordinates"}:
            continue
        if isinstance(v, str) and len(v) > 2048:
            continue
        cleaned[k] = v
    doc.metadata = cleaned
    return doc


# ============================================================
# HELPERS
# ============================================================
def get_embedding_dimension(embeddings) -> int:
    vec = embeddings.embed_query("dimension probe")
    return len(vec)


def batched(items: List[Document], n: int) -> Iterable[List[Document]]:
    for i in range(0, len(items), n):
        yield items[i : i + n]


def text_quality(text: str) -> Tuple[int, int, float]:
    t = (text or "").strip()
    if not t:
        return 0, 0, 0.0
    char_len = len(t)
    words = re.findall(r"[A-Za-z0-9]+", t)
    word_count = len(words)
    alpha = sum(1 for c in t if c.isalpha())
    alpha_ratio = alpha / max(1, char_len)
    return char_len, word_count, alpha_ratio


def should_ocr_extracted_text(
    extracted: str, *, min_chars: int, min_words: int, min_alpha_ratio: float
) -> bool:
    char_len, word_count, alpha_ratio = text_quality(extracted)
    if char_len < min_chars:
        return True
    if word_count < min_words:
        return True
    if alpha_ratio < min_alpha_ratio:
        return True
    return False


# ============================================================
# IMAGE PRE-PROCESSING
# ============================================================
def preprocess_for_ocr(image: Image.Image) -> Image.Image:
    """
    Sharpen and boost contrast before OCR.
    Helps on low-res scans, faded text, and photos of printed pages.
    """
    image = image.convert("RGB")
    image = image.filter(ImageFilter.SHARPEN)
    enhancer = ImageEnhance.Contrast(image)
    image = enhancer.enhance(1.4)
    return image


def render_page_image(fitz_page, zoom: float = 3.0) -> Image.Image:
    """Render a full PDF page to a PIL image."""
    mat = fitz.Matrix(zoom, zoom)
    pix = fitz_page.get_pixmap(matrix=mat, alpha=False)
    img = Image.open(io.BytesIO(pix.tobytes("png")))
    img = ImageOps.exif_transpose(img)
    return img.convert("RGB")


def render_region_image(fitz_page, bbox, zoom: float = 3.0) -> Image.Image:
    """Render a specific bbox region of a page (for columns, tables, etc.)."""
    clip = fitz.Rect(bbox)
    mat = fitz.Matrix(zoom, zoom)
    pix = fitz_page.get_pixmap(matrix=mat, clip=clip, alpha=False)
    img = Image.open(io.BytesIO(pix.tobytes("png")))
    return img.convert("RGB")


# ============================================================
# LAYOUT ANALYSIS — MULTI-COLUMN DETECTION
# ============================================================
def detect_columns(fitz_page) -> List[fitz.Rect]:
    """
    Detect multi-column layout by clustering word x-midpoints into a
    horizontal histogram and finding low-density gaps between columns.

    Returns a list of column Rect objects sorted left-to-right.
    Returns a single full-page Rect for single-column pages.
    """
    try:
        words = fitz_page.get_text("words")
    except Exception:
        return [fitz_page.rect]

    if not words:
        return [fitz_page.rect]

    page_width = fitz_page.rect.width
    mid_xs = [((w[0] + w[2]) / 2) for w in words]

    # 10-bucket horizontal histogram
    bucket_size = page_width / 10
    buckets = [0] * 10
    for mx in mid_xs:
        b = min(int(mx / bucket_size), 9)
        buckets[b] += 1

    max_count = max(buckets) if buckets else 1
    threshold = max_count * 0.15  # buckets below 15% of peak = gap

    valleys = [i for i, cnt in enumerate(buckets) if cnt <= threshold]

    # Group contiguous valley buckets into gap ranges
    gap_ranges = []
    if valleys:
        start = valleys[0]
        prev = valleys[0]
        for v in valleys[1:]:
            if v == prev + 1:
                prev = v
            else:
                gap_ranges.append((start * bucket_size, (prev + 1) * bucket_size))
                start = v
                prev = v
        gap_ranges.append((start * bucket_size, (prev + 1) * bucket_size))

    # Only count interior gaps (not page margin gaps)
    interior_gaps = [
        g for g in gap_ranges
        if g[0] > page_width * 0.1 and g[1] < page_width * 0.9
    ]

    if not interior_gaps:
        return [fitz_page.rect]

    # Build column rects from gap midpoints
    page_y0 = fitz_page.rect.y0
    page_y1 = fitz_page.rect.y1
    boundaries = [fitz_page.rect.x0]
    for gx0, gx1 in sorted(interior_gaps):
        boundaries.append((gx0 + gx1) / 2)
    boundaries.append(fitz_page.rect.x1)

    columns = []
    for i in range(len(boundaries) - 1):
        columns.append(fitz.Rect(boundaries[i], page_y0, boundaries[i + 1], page_y1))

    logger.info(f"    Detected {len(columns)} column(s)")
    return columns


# ============================================================
# TEXT EXTRACTION
# ============================================================
def extract_text_from_page(plumber_page, fitz_page) -> str:
    """
    Primary extraction: pdfplumber first (better spacing/table awareness),
    fall back to fitz raw text.
    """
    text = ""
    try:
        text = plumber_page.extract_text(x_tolerance=3, y_tolerance=3) or ""
    except Exception:
        text = ""

    if not text.strip():
        try:
            text = fitz_page.get_text("text") or ""
        except Exception:
            text = ""

    return clean_text(text)


def extract_columns_text(fitz_page) -> str:
    """
    For multi-column pages: detect columns, extract each column's text
    separately, then join. Prevents words from adjacent columns being
    interleaved by a naive left-to-right read order.
    """
    columns = detect_columns(fitz_page)

    if len(columns) == 1:
        try:
            return clean_text(fitz_page.get_text("text"))
        except Exception:
            return ""

    column_texts = []
    for col_rect in columns:
        try:
            col_text = fitz_page.get_text("text", clip=col_rect)
            col_text = clean_text(col_text)
            if col_text.strip():
                column_texts.append(col_text)
        except Exception:
            continue

    return "\n\n".join(column_texts)


def ocr_page_image(image: Image.Image) -> str:
    image = preprocess_for_ocr(image)
    return generate_image_caption(
        image,
        prompt=(
            "Transcribe ALL visible text from this page exactly as written. "
            "Preserve paragraph breaks, bullet points, and numbered lists. "
            "For any tables or grids, output them as Markdown tables. "
            "On multi-column layouts, read the left column fully before the right. "
            "Do not summarize or paraphrase anything."
        ),
    )


def ocr_column(image: Image.Image, column_index: int, total_columns: int) -> str:
    """OCR a single column region with a column-aware prompt."""
    image = preprocess_for_ocr(image)
    return generate_image_caption(
        image,
        prompt=(
            f"This is column {column_index + 1} of {total_columns} on the page. "
            "Transcribe ALL text exactly as written, preserving lists, headings, "
            "and any table structure (use Markdown for tables). "
            "Do not summarize."
        ),
    )


# ============================================================
# PDF PROCESSING
# ============================================================
def process_pdf(filepath: str, filename: str) -> List[Document]:
    docs: List[Document] = []

    with fitz.open(filepath) as fdoc, pdfplumber.open(filepath) as pdoc:
        total_pages = len(fdoc)
        logger.info(f"Processing PDF: {filename} | Pages: {total_pages}")

        for page_idx in range(total_pages):
            fitz_page = fdoc.load_page(page_idx)
            plumber_page = pdoc.pages[page_idx] if page_idx < len(pdoc.pages) else None
            page_num = page_idx + 1

            # ── Word count check ──
            try:
                words_on_page = fitz_page.get_text("words") or []
                word_boxes = len(words_on_page)
            except Exception:
                word_boxes = 0

            # ── Primary extraction ──
            if plumber_page:
                extracted = extract_text_from_page(plumber_page, fitz_page)
            else:
                extracted = clean_text(fitz_page.get_text("text"))

            # ── Column-aware extraction (replaces naive extraction if better) ──
            if word_boxes >= 25 and not should_ocr_extracted_text(
                extracted, min_chars=180, min_words=40, min_alpha_ratio=0.12
            ):
                column_text = extract_columns_text(fitz_page)
                if len(column_text.strip()) > len(extracted.strip()):
                    extracted = column_text
                    logger.info(f"    p{page_num}: column-aware extraction used")

            # ── OCR decision ──
            ocr_needed = should_ocr_extracted_text(
                extracted, min_chars=180, min_words=40, min_alpha_ratio=0.12
            )
            if word_boxes < 25:
                ocr_needed = True

            page_type = "pdf"
            if ocr_needed:
                logger.info(f"    p{page_num}: OCR triggered (word_boxes={word_boxes})")
                columns = detect_columns(fitz_page)

                if len(columns) > 1:
                    # OCR each column separately — much more accurate on multi-column scans
                    column_texts = []
                    for col_i, col_rect in enumerate(columns):
                        try:
                            col_img = render_region_image(fitz_page, col_rect, zoom=3.5)
                            col_text = ocr_column(col_img, col_i, len(columns))
                            if col_text.strip():
                                column_texts.append(col_text)
                        except Exception as e:
                            logger.warning(f"Column OCR failed p{page_num} col{col_i}: {e}")
                    page_text = clean_text("\n\n".join(column_texts)) if column_texts else extracted
                else:
                    try:
                        img = render_page_image(fitz_page, zoom=3.5)
                        page_text = clean_text(ocr_page_image(img))
                    except Exception as e:
                        logger.warning(f"OCR failed {filename} p{page_num}: {e}")
                        page_text = extracted

                page_type = "pdf_ocr"
            else:
                page_text = extracted

            # ── Table extraction (native-text pages only; OCR'd pages capture
            #    tables in the OCR prompt already) ──
            if page_type == "pdf" and len(extracted.strip()) >= 200:
                page_text = _extract_and_append_tables(
                    fitz_page, page_text, filename, page_num
                )

            if page_text.strip():
                docs.append(
                    Document(
                        page_content=page_text,
                        metadata={
                            "source": filename,
                            "page": page_num,
                            "type": page_type,
                        },
                    )
                )

    return docs


def _extract_and_append_tables(
    fitz_page, existing_text: str, filename: str, page_num: int
) -> str:
    """
    Detect tables on a native-text page, OCR them using Pydantic Structured Outputs
    for perfect parsing, and append the clean data to the existing page text.
    """
    try:
        tables = fitz_page.find_tables()
    except Exception:
        return existing_text

    if not tables:
        return existing_text

    page_text = existing_text
    for t_i, table in enumerate(tables):
        tmp_path = None
        try:
            bbox = table.bbox
            # Skip noise (tiny bounding boxes)
            if (bbox[2] - bbox[0]) * (bbox[3] - bbox[1]) < 500:
                continue

            pix = fitz_page.get_pixmap(
                clip=fitz.Rect(bbox),
                matrix=fitz.Matrix(3, 3),
                alpha=False,
            )
            image = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
            image = preprocess_for_ocr(image)

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                tmp_path = tmp.name
                image.save(tmp_path)

            s3_key = f"tables/{filename}_p{page_num}_t{t_i}.png"
            if upload_file_to_s3(tmp_path, s3_key):
                
                # --- NEW: PYDANTIC STRUCTURED AI EXTRACTION ---
                logger.info(f"Running structured table extraction for table {t_i} on page {page_num}")
                chat = get_vision_client() 
                structured_chat = chat.with_structured_output(ExtractedTable)
                
                image_base64 = encode_image(image)
                msg = HumanMessage(
                    content=[
                        {"type": "text", "text": "Extract all rows from this curriculum or information table precisely. Keep all words and units together."},
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_base64}"}}
                    ]
                )
                
                try:
                    result = structured_chat.invoke([msg])
                    
                    # Convert the perfect JSON back into a clean, readable text block for Pinecone
                    # This builds rows like "Row: GE 1 | Readings in Philippine History | 3 | None | 3 | None | None"
                    clean_table_text = "\n".join([
                        f"Row: {r.column_1} | {r.column_2} | {r.column_3} | {r.column_4} | {r.column_5} | {r.column_6} | {r.column_7}".strip(' |')
                        for r in result.rows
                    ])
                    
                    if clean_table_text.strip():
                        page_text += f"\n\n[TABLE {t_i + 1}]\n{clean_table_text}\n"
                except Exception as e:
                    logger.warning(f"Structured table extraction failed p{page_num} t{t_i}: {e}")

        except Exception as e:
            logger.warning(f"Table image generation failed p{page_num} t{t_i}: {e}")
        finally:
            if tmp_path and os.path.exists(tmp_path):
                try:
                    os.remove(tmp_path)
                except Exception:
                    pass

    return page_text


# ============================================================
# DOCX LINEAR PROCESSING
# ============================================================
def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("unexpected parent type")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def get_heading_level(paragraph) -> Optional[int]:
    """Return heading level 1-9 if this paragraph is a heading style, else None."""
    style_name = (paragraph.style.name or "").lower()
    if style_name.startswith("heading"):
        try:
            return int(style_name.split()[-1])
        except (ValueError, IndexError):
            return 1
    return None


def get_images_from_paragraph(paragraph, doc_obj):
    images = []
    for run in paragraph.runs:
        if "w:drawing" in run.element.xml:
            drawing_xml = run.element.xml
            if 'r:embed="' in drawing_xml:
                try:
                    rId = drawing_xml.split('r:embed="')[1].split('"')[0]
                    if rId in doc_obj.part.related_parts:
                        image_part = doc_obj.part.related_parts[rId]
                        images.append(image_part.blob)
                except Exception:
                    pass
    return images


def _docx_table_to_markdown(table) -> str:
    """
    Convert a python-docx Table to a Markdown table string.
    Deduplicates merged cells (they share the same _tc object).
    """
    rows_data = []
    seen_cells = set()

    for row in table.rows:
        row_cells = []
        for cell in row.cells:
            cell_id = id(cell._tc)
            if cell_id in seen_cells:
                row_cells.append("")  # merged-cell placeholder
            else:
                seen_cells.add(cell_id)
                row_cells.append(cell.text.strip().replace("\n", " "))
        rows_data.append(row_cells)

    if not rows_data:
        return ""

    col_count = max(len(r) for r in rows_data)
    rows_data = [r + [""] * (col_count - len(r)) for r in rows_data]

    header = "| " + " | ".join(rows_data[0]) + " |"
    separator = "| " + " | ".join(["---"] * col_count) + " |"
    body_rows = ["| " + " | ".join(r) + " |" for r in rows_data[1:]]

    return "\n".join([header, separator] + body_rows)


def process_docx_linear(filepath: str, filename: str) -> List[Document]:
    """
    Read DOCX linearly, preserving document order.
    """
    logger.info(f"Linear DOCX processing: {filename}")
    try:
        doc = docx.Document(filepath)
    except Exception as e:
        logger.error(f"Failed to open DOCX {filename}: {e}")
        return []

    sections: List[dict] = []
    current_section: dict = {"heading": "", "lines": []}
    context_buffer: List[str] = []

    def flush_section():
        if current_section["lines"]:
            sections.append({
                "heading": current_section["heading"],
                "lines": list(current_section["lines"]),
            })

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            level = get_heading_level(block)

            if level is not None and text:
                flush_section()
                current_section = {
                    "heading": text,
                    "lines": [f"{'#' * level} {text}"],
                }
                context_buffer = [text]
                continue

            if text:
                current_section["lines"].append(text)
                context_buffer.append(text)
                if len(context_buffer) > 5:
                    context_buffer.pop(0)

            # Inline images
            for img_bytes in get_images_from_paragraph(block, doc):
                try:
                    image = Image.open(io.BytesIO(img_bytes)).convert("RGB")
                    image = preprocess_for_ocr(image)
                    context_str = " | ".join(context_buffer)
                    logger.info(f"  OCRing inline image in {filename}")
                    caption = generate_image_caption(
                        image,
                        prompt=(
                            f"Context from surrounding text: {context_str}. "
                            "Transcribe ALL text in this image exactly as written. "
                            "If it contains a table, output it as a Markdown table. "
                            "Do not summarize."
                        ),
                    )
                    if caption and caption.strip():
                        current_section["lines"].append(
                            f"\n[IMAGE — context: {context_str}]\n{caption}\n"
                        )
                except Exception as e:
                    logger.warning(f"Inline image OCR failed in {filename}: {e}")

        elif isinstance(block, Table):
            md_table = _docx_table_to_markdown(block)
            if md_table:
                context_str = " | ".join(context_buffer)
                current_section["lines"].append(
                    f"\n[TABLE — context: {context_str}]\n{md_table}\n"
                )

    flush_section()

    if not sections:
        return []

    documents = []
    for sec in sections:
        content = "\n\n".join(sec["lines"]).strip()
        if content:
            documents.append(
                Document(
                    page_content=content,
                    metadata={
                        "source": filename,
                        "type": "docx_section",
                        "section_heading": sec["heading"][:200],
                    },
                )
            )

    return documents


# ============================================================
# LOADERS
# ============================================================
def get_loader(filepath: str):
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".txt":
        return TextLoader(filepath)
    if ext == ".md":
        return UnstructuredMarkdownLoader(filepath)
    if ext == ".csv":
        return CSVLoader(filepath, encoding="utf-8")
    return None


# ============================================================
# INDEX MANAGEMENT
# ============================================================
def ensure_index_exists(pc: Pinecone, index_name: str, dimension: int):
    if index_name not in pc.list_indexes().names():
        logger.info(f"Creating index '{index_name}' with metric='dotproduct'...")
        pc.create_index(
            name=index_name,
            dimension=dimension,
            metric="dotproduct",
            spec=ServerlessSpec(cloud="aws", region="us-east-1"),
        )
        while not pc.describe_index(index_name).status["ready"]:
            time.sleep(1)


# ============================================================
# SHARED CHUNKING HELPER
# ============================================================
def _chunk_and_dedup(raw_docs: List[Document]) -> List[Document]:
    """Chunk a list of Documents with Cross-Page Stitching and Persistent Headers."""
    final_chunks: List[Document] = []
    seen_ids = set()

    current_header = "General College Information"
    current_source = None
    previous_tail = ""

    for doc in raw_docs:
        source = doc.metadata.get("source", "")
        
        # Reset trackers if we switch to a completely new PDF
        if source != current_source:
            current_header = "General College Information"
            current_source = source
            previous_tail = ""

        clean_metadata(doc)
        
        # 🚀 CROSS-PAGE STITCHING: Attach the previous page's ending to this page's beginning
        combined_content = doc.page_content
        if previous_tail:
            combined_content = f"{previous_tail}\n{combined_content}"
        
        # Save the last 40 words of THIS page to carry over to the NEXT page
        words = doc.page_content.split()
        previous_tail = " ".join(words[-40:]) if len(words) >= 40 else doc.page_content

        # 🚀 PERSISTENT HEADERS: Pass the running header in, get the updated one back out
        chunks, current_header = smart_chunking(
            combined_content, 
            running_header=current_header
        )

        for idx, text in enumerate(chunks):
            if is_low_value_chunk(text):
                continue

            cid = generate_chunk_id(
                doc.metadata.get("source", ""),
                int(doc.metadata.get("page", 0) or 0),
                idx,
                text,
            )
            if cid in seen_ids:
                continue
            seen_ids.add(cid)

            meta = dict(doc.metadata)
            meta["chunk_index"] = idx
            meta["chunk_id"] = cid
            final_chunks.append(Document(page_content=text, metadata=meta))

    return final_chunks


# ============================================================
# BUILD / APPEND
# ============================================================
def build_index(
    data_path: str = "data/",
    index_name: str = "rag-google-v1",
    embeddings=None,
):
    if embeddings is None:
        raise ValueError("Embeddings required")

    dim = get_embedding_dimension(embeddings)
    pc = Pinecone(api_key=PINECONE_API_KEY)
    ensure_index_exists(pc, index_name, dimension=dim)

    raw_docs: List[Document] = []
    logger.info("Gathering documents...")

    for fname in os.listdir(data_path):
        path = os.path.join(data_path, fname)
        if not os.path.isfile(path):
            continue
        lower = fname.lower()
        if lower.endswith(".pdf"):
            raw_docs.extend(process_pdf(path, fname))
        elif lower.endswith(".docx"):
            raw_docs.extend(process_docx_linear(path, fname))
        else:
            loader = get_loader(path)
            if loader:
                docs = loader.load()
                for d in docs:
                    d.metadata["source"] = fname
                    clean_metadata(d)
                raw_docs.extend(docs)

    logger.info("Chunking documents...")
    final_chunks = _chunk_and_dedup(raw_docs)
    logger.info(f"Total chunks after dedup: {len(final_chunks)}")

    logger.info(f"Training BM25 on {len(final_chunks)} chunks...")
    bm25 = BM25Encoder().default()
    bm25.fit([d.page_content for d in final_chunks])
    bm25.dump(BM25_VALUES_FILE)
    logger.info(f"BM25 saved to {BM25_VALUES_FILE}")

    logger.info(f"Upserting {len(final_chunks)} chunks...")
    index = pc.Index(index_name)
    retriever = PineconeHybridSearchRetriever(
        embeddings=embeddings,
        sparse_encoder=bm25,
        index=index,
    )
    for batch in batched(final_chunks, 100):
        retriever.add_texts(
            [d.page_content for d in batch],
            metadatas=[d.metadata for d in batch],
        )

    logger.info("Index build complete.")


def append_file_to_index(
    filepath: str,
    index_name: str = "rag-google-v1",
    real_name: Optional[str] = None,
    embeddings=None,
) -> int:
    if embeddings is None:
        raise ValueError("Embeddings required")

    filename = real_name if real_name else os.path.basename(filepath)
    dim = get_embedding_dimension(embeddings)
    pc = Pinecone(api_key=PINECONE_API_KEY)
    ensure_index_exists(pc, index_name, dimension=dim)

    if os.path.exists(BM25_VALUES_FILE):
        bm25 = BM25Encoder().load(BM25_VALUES_FILE)
    else:
        logger.warning("BM25 file not found — using default encoder.")
        bm25 = BM25Encoder().default()

    raw_docs: List[Document] = []
    lower = filename.lower()
    if lower.endswith(".pdf"):
        raw_docs = process_pdf(filepath, filename)
    elif lower.endswith(".docx"):
        raw_docs = process_docx_linear(filepath, filename)
    else:
        loader = get_loader(filepath)
        if not loader:
            raise ValueError(f"Unsupported file type: {filename}")
        raw_docs = loader.load()
        for d in raw_docs:
            d.metadata["source"] = filename
            clean_metadata(d)

    final_chunks = _chunk_and_dedup(raw_docs)

    if final_chunks:
        index = pc.Index(index_name)
        retriever = PineconeHybridSearchRetriever(
            embeddings=embeddings,
            sparse_encoder=bm25,
            index=index,
        )
        for batch in batched(final_chunks, 100):
            retriever.add_texts(
                [d.page_content for d in batch],
                metadatas=[d.metadata for d in batch],
            )

    logger.info(f"Appended {len(final_chunks)} chunks from {filename}")
    return len(final_chunks)


if __name__ == "__main__":
    from config import INDEX_NAME
    logging.basicConfig(level=logging.INFO)
    embs = get_local_embeddings()
    build_index(index_name=INDEX_NAME, embeddings=embs)